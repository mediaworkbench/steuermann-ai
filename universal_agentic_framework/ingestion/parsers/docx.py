"""DOCX document parser using python-docx."""

from pathlib import Path
from typing import Dict, Any

from docx import Document

from universal_agentic_framework.ingestion.parsers.base import DocumentParser


class DOCXParser(DocumentParser):
    """Parser for DOCX documents."""
    
    def parse(self, file_path: Path) -> str:
        """Extract text from DOCX.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        doc = Document(str(file_path))
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return "\n\n".join(text_parts)
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with title, author, subject, etc.
        """
        doc = Document(str(file_path))
        metadata = {}
        
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "comments": props.comments or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }
        
        # Clean empty values
        metadata = {k: v for k, v in metadata.items() if v}
        metadata["file_name"] = file_path.name
        metadata["file_extension"] = file_path.suffix
        metadata["paragraphs"] = len(doc.paragraphs)
        
        return metadata
    
    @classmethod
    def supports_extension(cls, extension: str) -> bool:
        """Check if extension is .docx."""
        return extension.lower() == ".docx"
